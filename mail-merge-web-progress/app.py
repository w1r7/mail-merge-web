import os
import re
import time
import uuid
import zipfile
import tempfile
from threading import Thread
from datetime import datetime

import pandas as pd
import pytz
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docxcompose.composer import Composer

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200MB uploads

# In-memory job store
JOBS = {}

MERGE_FIELDS = {
    "<<DATE>>": "DATE",
    "<<JOB #>>": "JOB #",
    "<<NAME>>": "NAME",
    "<<VEHICLE MAKE>>": "VEHICLE MAKE",
    "<<VEHICLE MODEL>>": "VEHICLE MODEL",
    "<<KCML SERIAL #>>": "KCML SERIAL #",
    "<<VOLTAGE USED>>": "VOLTAGE USED",
    "<<CAN OR FREQUENCY>>": "CAN OR FREQUENCY",
    "<<PPK>>": "PPK",
    "<<SAP FILE>>": "SAP FILE",
    "<<ECU SERIAL #>>": "ECU SERIAL #",
}

def today_yymmdd_brisbane():
    tz = pytz.timezone("Australia/Brisbane")
    return datetime.now(tz).strftime("%y%m%d")

def format_value(val):
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return ""
    if isinstance(val, float) and val.is_integer():
        return str(int(val))
    if isinstance(val, (pd.Timestamp, datetime)):
        return val.strftime("%Y-%m-%d")
    return str(val)

def replace_placeholders_everywhere(doc, data_row):
    replacements = {ph: format_value(data_row.get(col, "")) for ph, col in MERGE_FIELDS.items()}

    def process_paragraph(paragraph):
        runs = paragraph.runs
        if not runs:
            return
        full_text = ''.join(run.text for run in runs)
        for ph, val in replacements.items():
            if ph not in full_text:
                continue
            start = 0
            while True:
                idx = full_text.find(ph, start)
                if idx == -1:
                    break
                # find run indices covering placeholder
                char_count = 0
                start_run = end_run = None
                for i, run in enumerate(runs):
                    rl = len(run.text)
                    if rl == 0: 
                        continue
                    seg_start = char_count
                    seg_end = char_count + rl - 1
                    if start_run is None and seg_start <= idx <= seg_end:
                        start_run = i
                    if seg_start <= idx + len(ph) - 1 <= seg_end:
                        end_run = i
                        break
                    char_count += rl
                if start_run is not None and end_run is not None:
                    combined = ''.join(runs[i].text for i in range(start_run, end_run + 1))
                    replaced = combined.replace(ph, val)
                    for i in range(start_run, end_run + 1):
                        runs[i].text = ''
                    runs[start_run].text = replaced
                    full_text = ''.join(run.text for run in runs)
                    start = idx + len(val)
                else:
                    start = idx + len(ph)

    for para in doc.paragraphs:
        process_paragraph(para)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)
    # headers/footers too
    for section in doc.sections:
        for para in section.header.paragraphs:
            process_paragraph(para)
        for para in section.footer.paragraphs:
            process_paragraph(para)
        for tbl in section.header.tables:
            for r in tbl.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        process_paragraph(p)
        for tbl in section.footer.tables:
            for r in tbl.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        process_paragraph(p)

def merge_worker(job_id, excel_path, template_paths, row_start, row_end):
    job = JOBS[job_id]
    job.update(status='running', message='Reading Excel…', start_time=time.time(), completed=0, total=0)
    try:
        df = pd.read_excel(excel_path, skiprows=2)
        df.fillna("", inplace=True)

        if row_start < 4:
            raise ValueError("❌ Invalid range: You must choose rows 4 or higher. Row 3 is for headers.")
        max_row = df.shape[0] + 3
        if row_end > max_row:
            raise ValueError(f"❌ Invalid range: Excel only has data up to row {max_row}. You requested row {row_end}.")

        num_rows = row_end - row_start + 1
        job['total'] = num_rows * len(template_paths) + len(template_paths)

        yymmdd = today_yymmdd_brisbane()
        session_dir = tempfile.mkdtemp(prefix=f"merge_{job_id}_")
        outputs = []

        for idx_t, template_path in enumerate(template_paths, start=1):
            job['message'] = f"Processing template {idx_t}/{len(template_paths)}…"
            per_row_docs = []
            for excel_row in range(row_start, row_end + 1):
                row = df.iloc[excel_row - 4]
                doc = Document(template_path)
                replace_placeholders_everywhere(doc, row)
                out_path = os.path.join(session_dir, f"Row_{excel_row}.docx")
                doc.save(out_path)
                per_row_docs.append(out_path)
                job['completed'] += 1
                # ETA calc
                elapsed = time.time() - job['start_time']
                done = job['completed']
                total = max(1, job['total'])
                job['eta'] = int(elapsed / done * (total - done)) if done else None

            # combine preserving headers/footers
            base_doc = Document(per_row_docs[0])
            composer = Composer(base_doc)
            for p in per_row_docs[1:]:
                composer.append(Document(p))

            base_name = os.path.splitext(os.path.basename(template_path))[0]
            short_match = re.search(r"(PRF\s?\d{3}-\d)", base_name.upper())
            short_code = short_match.group(1).replace(" ", "") if short_match else "Template"

            final_name = f"{yymmdd}_row{row_start}-{row_end}_of-{short_code}.docx"
            final_path = os.path.join(session_dir, final_name)
            composer.save(final_path)
            outputs.append(final_path)
            job['completed'] += 1

        if len(outputs) > 1:
            zip_name = f"{yymmdd}_row{row_start}-{row_end}_of_{len(outputs)}_files.zip"
            zip_path = os.path.join(session_dir, zip_name)
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
                for f in outputs:
                    z.write(f, os.path.basename(f))
            job.update(status='done', message='✅ Done', result_path=zip_path, result_name=os.path.basename(zip_path))
        else:
            job.update(status='done', message='✅ Done', result_path=outputs[0], result_name=os.path.basename(outputs[0]))
    except Exception as e:
        job.update(status='error', message='❌ Error', error=str(e))

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/start', methods=['POST'])
def start():
    try:
        row_start = int(request.form['row_start'])
        row_end = int(request.form['row_end'])
        if row_start < 4 or row_end < 4:
            return jsonify({'error': '❌ Invalid range: rows must be ≥ 4.'}), 400
        if row_start > row_end:
            return jsonify({'error': 'Row start cannot be greater than row end.'}), 400

        excel_file = request.files.get('excel')
        template_files = request.files.getlist('templates')
        if not excel_file or not template_files:
            return jsonify({'error': 'Please upload one Excel and at least one Word template.'}), 400

        work_dir = tempfile.mkdtemp(prefix='uploads_')
        excel_path = os.path.join(work_dir, excel_file.filename)
        excel_file.save(excel_path)
        template_paths = []
        for f in template_files:
            t_path = os.path.join(work_dir, f.filename)
            f.save(t_path)
            template_paths.append(t_path)

        job_id = uuid.uuid4().hex[:12]
        JOBS[job_id] = {'status': 'queued', 'message': 'Queued…', 'total': 0, 'completed': 0, 'eta': None,
                        'result_path': None, 'result_name': None, 'start_time': None}
        Thread(target=merge_worker, args=(job_id, excel_path, template_paths, row_start, row_end), daemon=True).start()
        return jsonify({'job_id': job_id})
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/progress/<job_id>', methods=['GET'])
def progress(job_id):
    job = JOBS.get(job_id)
    if not job:
        return jsonify({'error': 'Invalid job id'}), 404
    total = max(1, job.get('total', 1))
    completed = job.get('completed', 0)
    pct = int(completed / total * 100)
    return jsonify({
        'status': job['status'],
        'message': job['message'],
        'progress': pct,
        'eta_seconds': job.get('eta'),
        'result_ready': job['status'] == 'done',
        'result_name': job.get('result_name')
    })

@app.route('/download/<job_id>', methods=['GET'])
def download(job_id):
    job = JOBS.get(job_id)
    if not job or job['status'] != 'done' or not job.get('result_path'):
        return "Result not ready", 400
    return send_file(job['result_path'], as_attachment=True, download_name=job['result_name'])

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)
